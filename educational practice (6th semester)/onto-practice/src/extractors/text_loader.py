"""Universal text loader: PDF (via OCR) / DOCX / TXT → list of page-sized chunks.

Returns one string per logical page. For PDF this is one entry per scanned page;
for DOCX / TXT we slice the document into ~3 KB chunks so the per-chunk prompt
stays small (free-tier text models choke on long context).
"""

from __future__ import annotations

import logging
import tempfile
from pathlib import Path

from docx import Document

from src.extractors.ocr import env_dpi, ocr_pages
from src.extractors.pdf_render import render_pages

log = logging.getLogger(__name__)

CHUNK_TARGET_CHARS = 3000


def load_pages(path: Path) -> list[str]:
    """Return a list of page/chunk text strings for any supported file type."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _chunk_text(_read_docx(path))
    if suffix == ".txt":
        return _chunk_text(path.read_text(encoding="utf-8", errors="replace"))
    raise ValueError(f"Unsupported file type: {suffix} (expected .pdf / .docx / .txt)")


def _load_pdf(path: Path) -> list[str]:
    with tempfile.TemporaryDirectory() as tmp:
        dpi = env_dpi()
        png_paths = render_pages(path, Path(tmp), dpi=dpi)
        log.info("rendered %d PDF pages at dpi=%d", len(png_paths), dpi)
        return ocr_pages(png_paths)


def _read_docx(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _chunk_text(text: str, target: int = CHUNK_TARGET_CHARS) -> list[str]:
    """Split text into chunks of ~target chars, preferring paragraph boundaries."""
    text = text.strip()
    if not text:
        return []
    if len(text) <= target:
        return [text]
    chunks: list[str] = []
    cur: list[str] = []
    cur_len = 0
    for para in text.split("\n"):
        line_len = len(para) + 1
        if cur and cur_len + line_len > target:
            chunks.append("\n".join(cur))
            cur, cur_len = [], 0
        cur.append(para)
        cur_len += line_len
    if cur:
        chunks.append("\n".join(cur))
    return chunks
